from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
import requests
import xml.etree.ElementTree as ET
from datetime import datetime
from docx import Document
from docx.shared import Inches, Mm
from docx.enum.section import WD_ORIENT
import PyPDF2
import google.generativeai as genai
import openai
import os
import tempfile
from werkzeug.utils import secure_filename
import re
from docx.shared import RGBColor

app = Flask(__name__, static_folder='.')
CORS(app)

# API 설정
OC = "climsneys85"  # 이메일 ID
search_url = "http://www.law.go.kr/DRF/lawSearch.do"
detail_url = "http://www.law.go.kr/DRF/lawService.do"

# 광역지자체 코드 및 이름
metropolitan_govs = {
    '6110000': '서울특별시',
    '6260000': '부산광역시',
    '6270000': '대구광역시',
    '6280000': '인천광역시',
    '6290000': '광주광역시',
    '6300000': '대전광역시',
    '5690000': '세종특별자치시',
    '6310000': '울산광역시',
    '6410000': '경기도',
    '6530000': '강원특별자치도',
    '6430000': '충청북도',
    '6440000': '충청남도',
    '6540000': '전북특별자치도',
    '6460000': '전라남도',
    '6470000': '경상북도',
    '6480000': '경상남도',
    '6500000': '제주특별자치도'
}

# 임시 파일 저장 경로
UPLOAD_FOLDER = tempfile.gettempdir()
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/')
def index():
    return send_from_directory('.', 'index.html')

@app.route('/<path:path>')
def static_files(path):
    return send_from_directory('.', path)

def get_ordinance_detail(ordinance_id):
    params = {
        'OC': OC,
        'target': 'ordin',
        'ID': ordinance_id,
        'type': 'XML'
    }
    try:
        response = requests.get(detail_url, params=params, timeout=60)
        root = ET.fromstring(response.text)
        articles = []
        for article in root.findall('.//조'):
            content = article.find('조내용').text if article.find('조내용') is not None else ""
            if content:
                content = content.replace('<![CDATA[', '').replace(']]>', '')
                content = content.replace('<p>', '').replace('</p>', '\n')
                content = content.replace('<br/>', '\n')
                content = content.replace('<br>', '\n')
                content = content.replace('&nbsp;', ' ')
                content = content.strip()
            if content:
                articles.append(content)
        return articles
    except Exception:
        return []

@app.route('/api/search', methods=['POST'])
def search():
    try:
        data = request.get_json()
        if not data or 'query' not in data:
            return jsonify({'error': '검색어가 필요합니다.'}), 400

        query = data['query'].strip()
        if not query:
            return jsonify({'error': '검색어가 비어있습니다.'}), 400

        results = []
        total_count = 0

        # 각 광역지자체별로 검색
        for org_code, metro_name in metropolitan_govs.items():
            try:
                params = {
                    'OC': OC,
                    'target': 'ordin',
                    'type': 'XML',
                    'query': query,
                    'display': 100,
                    'search': 1,  # 제목만 검색
                    'sort': 'ddes',
                    'page': 1,
                    'org': org_code
                }
                
                response = requests.get(search_url, params=params, timeout=60)
                response.raise_for_status()  # HTTP 오류 체크
                
                root = ET.fromstring(response.text)
                total_laws = len(root.findall('.//law'))
                
                if total_laws > 0:
                    for law in root.findall('.//law'):
                        ordinance_name = law.find('자치법규명').text if law.find('자치법규명') is not None else ""
                        ordinance_id = law.find('자치법규ID').text if law.find('자치법규ID') is not None else None
                        기관명 = law.find('지자체기관명').text if law.find('지자체기관명') is not None else ""
                        
                        if 기관명 != metro_name:
                            continue  # 본청이 아니면 건너뜀
                            
                        # 검색어 매칭 로직
                        search_terms = [term.lower() for term in query.split() if term.strip()]
                        ordinance_name_clean = ordinance_name.replace(' ', '').lower()
                        if not all(term in ordinance_name_clean for term in search_terms):
                            continue
                            
                        total_count += 1
                        articles = get_ordinance_detail(ordinance_id)
                        
                        results.append({
                            'name': ordinance_name,
                            'content': '\n'.join(articles) if articles else '(조문 없음)',
                            'metro': metro_name
                        })
                        
            except requests.RequestException as e:
                print(f"API 요청 오류 ({metro_name}): {str(e)}")
                continue
            except ET.ParseError as e:
                print(f"XML 파싱 오류 ({metro_name}): {str(e)}")
                continue
            except Exception as e:
                print(f"예상치 못한 오류 ({metro_name}): {str(e)}")
                continue

        return jsonify({
            'results': results,
            'total': total_count
        })

    except Exception as e:
        print(f"검색 처리 중 오류 발생: {str(e)}")
        return jsonify({'error': f'검색 처리 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/save', methods=['POST'])
def save():
    try:
        data = request.get_json()
        if not data or 'query' not in data:
            return jsonify({'error': '검색어가 필요합니다.'}), 400

        query = data['query'].strip()
        if not query:
            return jsonify({'error': '검색어가 비어있습니다.'}), 400

        # 검색 결과 수집
        results = []
        total_count = 0

        # 각 광역지자체별로 검색
        for org_code, metro_name in metropolitan_govs.items():
            try:
                params = {
                    'OC': OC,
                    'target': 'ordin',
                    'type': 'XML',
                    'query': query,
                    'display': 100,
                    'search': 1,  # 제목만 검색
                    'sort': 'ddes',
                    'page': 1,
                    'org': org_code
                }
                
                response = requests.get(search_url, params=params, timeout=60)
                response.raise_for_status()
                
                root = ET.fromstring(response.text)
                for law in root.findall('.//law'):
                    ordinance_name = law.find('자치법규명').text if law.find('자치법규명') is not None else ""
                    ordinance_id = law.find('자치법규ID').text if law.find('자치법규ID') is not None else None
                    기관명 = law.find('지자체기관명').text if law.find('지자체기관명') is not None else ""
                    
                    if 기관명 != metro_name:
                        continue  # 본청이 아니면 건너뜀
                        
                    # 검색어 매칭 로직
                    search_terms = [term.lower() for term in query.split() if term.strip()]
                    ordinance_name_clean = ordinance_name.replace(' ', '').lower()
                    if not all(term in ordinance_name_clean for term in search_terms):
                        continue
                        
                    total_count += 1
                    articles = get_ordinance_detail(ordinance_id)
                    results.append({
                        'name': ordinance_name,
                        'content': articles,
                        'metro': metro_name
                    })
                    
            except Exception as e:
                print(f"검색 중 오류 발생 ({metro_name}): {str(e)}")
                continue

        if not results:
            return jsonify({'error': '검색 결과가 없습니다.'}), 404

        # Word 문서 생성
        doc = Document()
        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(420)
        section.page_height = Mm(297)

        # 제목 추가
        doc.add_heading('조례 검색 결과', level=1)
        doc.add_paragraph(f'검색어: {query}')
        doc.add_paragraph(f'총 {total_count}건의 조례가 검색되었습니다.\n')

        # 조례를 3개씩 그룹화
        for i in range(0, len(results), 3):
            # 현재 페이지의 조례들
            current_laws = results[i:i+3]
            # 3개 미만이면 빈 값으로 채움
            while len(current_laws) < 3:
                current_laws.append({'name': '', 'content': [], 'metro': ''})

            # 표 생성 (1행, 3열 고정)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            table.autofit = True

            # 각 셀에 조례 내용 추가
            for idx, law in enumerate(current_laws):
                cell = table.cell(0, idx)
                paragraph = cell.paragraphs[0]
                
                if law['name']:
                    # 조례명 추가
                    run = paragraph.add_run(f"{law['metro']}\n{law['name']}\n")
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)  # 빨간색
                    
                    # 조문 내용 추가
                    if law['content']:
                        paragraph.add_run('\n'.join(law['content']))
                    else:
                        paragraph.add_run('(조문 없음)')

            # 마지막 페이지가 아니면 페이지 나누기 추가
            if i + 3 < len(results):
                doc.add_page_break()
                section = doc.sections[-1]
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Mm(420)
                section.page_height = Mm(297)

        # 임시 파일로 저장
        temp_docx = os.path.join(app.config['UPLOAD_FOLDER'], 'search_results.docx')
        doc.save(temp_docx)

        return send_file(
            temp_docx,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'조례_검색결과_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )

    except Exception as e:
        print(f"Word 문서 저장 중 오류 발생: {str(e)}")
        return jsonify({'error': f'Word 문서 저장 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/upload', methods=['POST'])
def upload():
    try:
        if 'pdf' not in request.files:
            return jsonify({'error': 'PDF 파일이 없습니다.'}), 400
            
        file = request.files['pdf']
        if file.filename == '':
            return jsonify({'error': '선택된 파일이 없습니다.'}), 400
            
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'error': 'PDF 파일만 업로드 가능합니다.'}), 400

        # 파일 저장
        filename = 'uploaded_pdf.pdf'
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # 파일이 제대로 저장되었는지 확인
        if not os.path.exists(filepath):
            return jsonify({'error': '파일 저장에 실패했습니다.'}), 500
            
        # PDF 파일 읽기 테스트
        try:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if len(reader.pages) == 0:
                    return jsonify({'error': 'PDF 파일이 비어있습니다.'}), 400
        except Exception as e:
            return jsonify({'error': f'PDF 파일 읽기 실패: {str(e)}'}), 400

        return jsonify({'message': 'PDF 파일이 성공적으로 업로드되었습니다.'})
        
    except Exception as e:
        print(f"PDF 업로드 중 오류 발생: {str(e)}")
        return jsonify({'error': f'PDF 업로드 중 오류가 발생했습니다: {str(e)}'}), 500

def extract_pdf_text(pdf_path):
    try:
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ''
            for page in reader.pages:
                text += page.extract_text() + '\n'
            return text
    except Exception as e:
        print(f"PDF 텍스트 추출 중 오류 발생: {str(e)}")
        return None

@app.route('/api/compare', methods=['POST'])
def compare():
    try:
        debug_logs = []
        if 'pdf' not in request.files:
            return jsonify({'error': 'PDF 파일이 없습니다.'}), 400

        pdf_file = request.files['pdf']
        if pdf_file.filename == '':
            return jsonify({'error': '선택된 파일이 없습니다.'}), 400

        if not pdf_file.filename.endswith('.pdf'):
            return jsonify({'error': 'PDF 파일만 업로드 가능합니다.'}), 400

        # PDF 파일 저장
        filename = secure_filename(pdf_file.filename)
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        pdf_file.save(pdf_path)

        # 검색 결과 수집
        query = request.form.get('query', '').strip()
        if not query:
            return jsonify({'error': '검색어가 필요합니다.'}), 400

        results = []
        total_count = 0

        # 각 광역지자체별로 검색
        for org_code, metro_name in metropolitan_govs.items():
            try:
                params = {
                    'OC': OC,
                    'target': 'ordin',
                    'type': 'XML',
                    'query': query,
                    'display': 100,
                    'search': 1,  # 제목만 검색
                    'sort': 'ddes',
                    'page': 1,
                    'org': org_code
                }
                
                response = requests.get(search_url, params=params, timeout=60)
                response.raise_for_status()
                
                root = ET.fromstring(response.text)
                for law in root.findall('.//law'):
                    ordinance_name = law.find('자치법규명').text if law.find('자치법규명') is not None else ""
                    ordinance_id = law.find('자치법규ID').text if law.find('자치법규ID') is not None else None
                    기관명 = law.find('지자체기관명').text if law.find('지자체기관명') is not None else ""
                    
                    if 기관명 != metro_name:
                        continue  # 본청이 아니면 건너뜀
                        
                    # 검색어 매칭 로직
                    search_terms = [term.lower() for term in query.split() if term.strip()]
                    ordinance_name_clean = ordinance_name.replace(' ', '').lower()
                    if not all(term in ordinance_name_clean for term in search_terms):
                        continue
                        
                    total_count += 1
                    articles = get_ordinance_detail(ordinance_id)
                    results.append({
                        'name': ordinance_name,
                        'content': articles,
                        'metro': metro_name
                    })
                    
            except Exception as e:
                debug_logs.append(f"검색 중 오류 발생 ({metro_name}): {str(e)}")
                continue

        # PDF 텍스트 추출
        pdf_text = extract_pdf_text(pdf_path)

        # API 키 확인
        gemini_api_key = request.form.get('geminiApiKey', '').strip()
        openai_api_key = request.form.get('openaiApiKey', '').strip()

        if not gemini_api_key and not openai_api_key:
            return jsonify({'error': 'API 키를 하나 이상 입력해주세요.'}), 400

        analysis_results = []
        is_first_ordinance = not results

        # Gemini API 분석
        if gemini_api_key:
            try:
                genai.configure(api_key=gemini_api_key)
                model = genai.GenerativeModel('gemini-1.5-flash')
                prompt = create_analysis_prompt(pdf_text, results, is_first_ordinance)
                debug_logs.append(f"[DEBUG] Gemini 프롬프트 길이: {len(prompt)}")
                response = model.generate_content(prompt)
                debug_logs.append(f"[DEBUG] Gemini 응답: {getattr(response, 'text', None)}")
                if response and hasattr(response, 'text') and response.text:
                    analysis_results.append({
                        'model': 'Gemini',
                        'content': response.text
                    })
                else:
                    analysis_results.append({
                        'model': 'Gemini',
                        'error': 'Gemini API 응답이 비어있음 또는 None입니다.'
                    })
            except Exception as e:
                debug_logs.append(f"Gemini API 오류: {str(e)}")
                analysis_results.append({
                    'model': 'Gemini',
                    'error': str(e)
                })

        # OpenAI API 분석
        if openai_api_key:
            try:
                openai.api_key = openai_api_key
                prompt = create_analysis_prompt(pdf_text, results, is_first_ordinance)
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "당신은 법률 전문가입니다. 조례 분석과 검토를 도와주세요."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=4000
                )
                if response.choices[0].message.content:
                    analysis_results.append({
                        'model': 'OpenAI',
                        'content': response.choices[0].message.content
                    })
            except Exception as e:
                print(f"OpenAI API 오류: {str(e)}")
                analysis_results.append({
                    'model': 'OpenAI',
                    'error': str(e)
                })

        if not analysis_results:
            return jsonify({'error': '분석 결과가 없습니다.'}), 500

        # 디버그 로그와 분석 결과만 반환
        return jsonify({
            'result': analysis_results,
            'debug': debug_logs
        })

    except Exception as e:
        return jsonify({'error': f'비교 분석 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/compare/download', methods=['POST'])
def compare_download():
    try:
        debug_logs = []
        if 'pdf' not in request.files:
            return jsonify({'error': 'PDF 파일이 없습니다.'}), 400

        pdf_file = request.files['pdf']
        if pdf_file.filename == '':
            return jsonify({'error': '선택된 파일이 없습니다.'}), 400

        if not pdf_file.filename.endswith('.pdf'):
            return jsonify({'error': 'PDF 파일만 업로드 가능합니다.'}), 400

        # PDF 파일 저장
        filename = secure_filename(pdf_file.filename)
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        pdf_file.save(pdf_path)

        # 검색 결과 수집
        query = request.form.get('query', '').strip()
        if not query:
            return jsonify({'error': '검색어가 필요합니다.'}), 400

        results = []
        total_count = 0

        # 각 광역지자체별로 검색
        for org_code, metro_name in metropolitan_govs.items():
            try:
                params = {
                    'OC': OC,
                    'target': 'ordin',
                    'type': 'XML',
                    'query': query,
                    'display': 100,
                    'search': 1,  # 제목만 검색
                    'sort': 'ddes',
                    'page': 1,
                    'org': org_code
                }
                
                response = requests.get(search_url, params=params, timeout=60)
                response.raise_for_status()
                
                root = ET.fromstring(response.text)
                for law in root.findall('.//law'):
                    ordinance_name = law.find('자치법규명').text if law.find('자치법규명') is not None else ""
                    ordinance_id = law.find('자치법규ID').text if law.find('자치법규ID') is not None else None
                    기관명 = law.find('지자체기관명').text if law.find('지자체기관명') is not None else ""
                    
                    if 기관명 != metro_name:
                        continue  # 본청이 아니면 건너뜀
                        
                    # 검색어 매칭 로직
                    search_terms = [term.lower() for term in query.split() if term.strip()]
                    ordinance_name_clean = ordinance_name.replace(' ', '').lower()
                    if not all(term in ordinance_name_clean for term in search_terms):
                        continue
                        
                    total_count += 1
                    articles = get_ordinance_detail(ordinance_id)
                    results.append({
                        'name': ordinance_name,
                        'content': articles,
                        'metro': metro_name
                    })
                    
            except Exception as e:
                debug_logs.append(f"검색 중 오류 발생 ({metro_name}): {str(e)}")
                continue

        # PDF 텍스트 추출
        pdf_text = extract_pdf_text(pdf_path)

        # API 키 확인
        gemini_api_key = request.form.get('geminiApiKey', '').strip()
        openai_api_key = request.form.get('openaiApiKey', '').strip()

        if not gemini_api_key and not openai_api_key:
            return jsonify({'error': 'API 키를 하나 이상 입력해주세요.'}), 400

        analysis_results = []
        is_first_ordinance = not results

        # Gemini API 분석
        if gemini_api_key:
            try:
                genai.configure(api_key=gemini_api_key)
                model = genai.GenerativeModel('gemini-1.5-flash')
                prompt = create_analysis_prompt(pdf_text, results, is_first_ordinance)
                debug_logs.append(f"[DEBUG] Gemini 프롬프트 길이: {len(prompt)}")
                response = model.generate_content(prompt)
                debug_logs.append(f"[DEBUG] Gemini 응답: {getattr(response, 'text', None)}")
                if response and hasattr(response, 'text') and response.text:
                    analysis_results.append({
                        'model': 'Gemini',
                        'content': response.text
                    })
                else:
                    analysis_results.append({
                        'model': 'Gemini',
                        'error': 'Gemini API 응답이 비어있음 또는 None입니다.'
                    })
            except Exception as e:
                debug_logs.append(f"Gemini API 오류: {str(e)}")
                analysis_results.append({
                    'model': 'Gemini',
                    'error': str(e)
                })

        # OpenAI API 분석
        if openai_api_key:
            try:
                openai.api_key = openai_api_key
                prompt = create_analysis_prompt(pdf_text, results, is_first_ordinance)
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "당신은 법률 전문가입니다. 조례 분석과 검토를 도와주세요."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=4000
                )
                if response.choices[0].message.content:
                    analysis_results.append({
                        'model': 'OpenAI',
                        'content': response.choices[0].message.content
                    })
            except Exception as e:
                print(f"OpenAI API 오류: {str(e)}")
                analysis_results.append({
                    'model': 'OpenAI',
                    'error': str(e)
                })

        if not analysis_results:
            return jsonify({'error': '분석 결과가 없습니다.'}), 500

        doc = create_comparison_document(pdf_text, results, analysis_results, debug_logs)
        temp_docx = os.path.join(app.config['UPLOAD_FOLDER'], 'comparison_results.docx')
        doc.save(temp_docx)
        return send_file(
            temp_docx,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'조례_비교분석_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )

    except Exception as e:
        return jsonify({'error': f'비교 분석(워드 저장) 중 오류가 발생했습니다: {str(e)}'}), 500

def create_analysis_prompt(pdf_text, search_results, is_first_ordinance=False):
    prompt = (
        "아래는 내가 업로드한 조례 PDF의 전체 내용이야.\n"
        "---\n"
        f"{pdf_text}\n"
        "---\n"
    )
    
    if is_first_ordinance:
        prompt += (
            "※ 참고: 이 조례는 17개 시도 중 최초로 제정되는 조례로, 타시도 조례가 존재하지 않습니다.\n"
            "타시도 조례가 없는 상황에서, 아래 기준에 따라 조례의 적정성, 상위법령과의 관계, 실무적 검토 포인트 등을 중심으로 분석해줘.\n"
        )
    else:
        prompt += "그리고 아래는 타시도 조례명과 각 조문 내용이야.\n"
        for result in search_results:
            prompt += f"조례명: {result['name']}\n"
            for idx, article in enumerate(result['content']):
                prompt += f"제{idx+1}조: {article}\n"
    
    prompt += (
        "---\n"
        "아래 기준에 따라 분석해줘. 반드시 한글로 답변해줘.\n"
        "1. [비교분석 요약표(조문별)]\n"
        "- 표의 컬럼: 조문(내 조례), 주요 내용, 타 시도 유사 조항, 동일 여부, 차이 및 내 조례 특징, 추천 조문\n"
        "- 반드시 내 조례(PDF로 업로드한 조례)의 조문만을 기준으로, 각 조문별로 타 시도 조례와 비교해 표로 정리(내 조례에 없는 조문은 비교하지 말 것)\n"
        "- '추천 조문' 칸에는 타 시도 조례와 비교해 무난하게 생각되는 조문 예시를 한글로 작성\n\n"
        "2. [내 조례의 차별점 요약] (별도 소제목)\n"
        "- 타 시도 조례와 비교해 독특하거나 구조적으로 다른 점, 내 조례만의 관리/운영 방식 등 요약\n\n"
        "3. [검토 시 유의사항] (별도 소제목)\n"
        "각 항목마다 일반인도 이해할 수 있도록 쉬운 말로 부연설명도 함께 작성해줘.\n"
        "다음 원칙들을 기준으로 검토해줘:\n"
        "a) 소관사무의 원칙\n"
        "- 지방자치단체의 자치사무와 법령에 의해 위임된 단체위임사무에 대해서만 제정 가능한지\n"
        "- 국가사무가 지방자치단체의 장에게 위임된 기관위임사무인 경우 조례 제정이 적절한지\n"
        "- 사무의 성격이 전국적으로 통일적 처리를 요구하는지 여부 검토\n\n"
        "b) 법률 유보의 원칙\n"
        "- 주민의 권리를 제한하거나 의무를 부과하는 내용이 있는지\n"
        "- 상위 법령에서 위임받지 않은 권한을 행사하는지\n"
        "- 상위 법령의 위임 범위를 초과하는지\n\n"
        "c) [법령우위의 원칙 위반 여부]\n"
        "- 조례가 상위 법령의 내용과 직접적으로 충돌하거나 위배되는지\n"
        "- 상위 법령의 취지나 목적을 해치는지\n"
        "- 상위 법령이 금지하는 행위를 허용하거나, 의무화하는 행위를 면제하는지\n\n"
        "4. [실무적 검토 포인트]\n"
        "- 조례의 집행 과정에서 발생할 수 있는 문제점\n"
        "- 개선이 필요한 부분과 그 방향성\n"

    )
    return prompt

def create_comparison_document(pdf_text, search_results, analysis_results, debug_logs=None):
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(420)
    section.page_height = Mm(297)

    # 제목 추가
    doc.add_heading('조례 비교 분석 결과', level=1)
    doc.add_paragraph(f'분석 일시: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')

    # 디버그 로그가 있으면 문서에 추가
    if debug_logs:
        doc.add_heading('디버그 로그', level=2)
        for log in debug_logs:
            doc.add_paragraph(log)

    # 각 API 분석 결과 추가
    for result in analysis_results:
        if 'error' in result:
            doc.add_paragraph(f"{result['model']} API 오류: {result['error']}")
            continue

        # 분석 결과 텍스트 정리
        content = result['content']
        
        # 1. 비교분석 요약표 추출 및 표로 변환
        table_pattern = re.compile(r'(\|.+\|\n)+')
        table_match = table_pattern.search(content)
        if table_match:
            table_text = table_match.group()
            rows = [row.strip() for row in table_text.strip().split('\n') if row.strip()]
            # 마크다운 구분선(---) 제거
            rows = [row for row in rows if not set(row.replace('|','').strip()) <= set('-')]
            table_data = [[cell.strip().replace('**','') for cell in row.split('|')[1:-1]] for row in rows]
            
            # 표 생성
            table = doc.add_table(rows=1, cols=len(table_data[0]))
            table.style = 'Table Grid'
            table.autofit = True
            
            # 헤더 추가
            for i, cell in enumerate(table.rows[0].cells):
                cell.text = table_data[0][i]
                if '동일 여부' in cell.text:
                    cell.width = Mm(20)
                if '추천 조문' in cell.text:
                    cell.width = Mm(80)
            
            # 데이터 추가
            for row in table_data[1:]:
                cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    cells[i].text = row[i]
            
            doc.add_paragraph('')
            # 표 이후 텍스트만 남기기 위해 결과에서 표 부분 제거
            content = content.replace(table_text, '')

        # 2. 차별점 요약, 3. 검토시 유의사항 등 나머지 텍스트(마크다운 기호 제거)
        clean_text = re.sub(r'[#*`>\-]+', '', content)
        
        # '3. 검토 시 유의사항'과 '4.' 블록 모두에서 상위법령 후보 추출
        def extract_law_sections(text):
            # '3. 검토 시 유의사항' ~ '4.' 또는 'd)' 또는 끝까지
            m1 = re.search(r'3[.)]\s*검토 시 유의사항[\s\S]+?(?=\n4[.)]|\nd[.)]|$)', text)
            section1 = m1.group(0) if m1 else ''
            # 'd) [실무적 검토 포인트]' ~ '4.' 또는 끝까지
            m2 = re.search(r'd[.)]\s*\[실무적 검토 포인트\][\s\S]+?(?=\n4[.)]|$)', text)
            section2 = m2.group(0) if m2 else ''
            return section1 + '\n' + section2

        law_section = extract_law_sections(clean_text)
        
        # 상위법령 후보 추출
        upper_law_candidates = set()
        law_pattern = re.compile(r'([가-힣·\s]{2,}?(법|시행령|시행규칙))')
        for m in law_pattern.finditer(law_section):
            law_name = m.group(1).strip('「」[]()<>"" .,;:!?~-')
            if is_valid_law_name(law_name):
                upper_law_candidates.add(law_name)

        # 나머지 분석 결과 추가 (중복 문단 제거)
        added_paragraphs = set()
        for line in clean_text.split('\n'):
            line_strip = line.strip()
            if line_strip and line_strip not in added_paragraphs:
                doc.add_paragraph(line_strip)
                added_paragraphs.add(line_strip)

        # 상위법령 위반 여부 검토 (문서 마지막에 추가)
        if upper_law_candidates:
            doc.add_page_break()  # 새로운 페이지 시작
            doc.add_heading('상위법령 위반 여부 검토', level=1)
            for upper_law_name in upper_law_candidates:
                try:
                    print(f"[DEBUG] 상위법령명: {upper_law_name}")
                    # 1. lawSearch로 현행 법령ID 및 법령명한글 얻기
                    search_params = {
                        'OC': OC,
                        'target': 'law',
                        'type': 'XML',
                        'query': upper_law_name
                    }
                    search_resp = requests.get(search_url, params=search_params, timeout=60)
                    print(f"[DEBUG] lawSearch status: {search_resp.status_code}")
                    search_root = ET.fromstring(search_resp.text)
                    law_id = None
                    law_name_kor = None
                    for law in search_root.findall('.//law'):
                        if law.find('현행연혁코드') is not None and law.find('현행연혁코드').text == '현행':
                            law_id = law.find('법령ID').text if law.find('법령ID') is not None else None
                            law_name_kor = law.find('법령명한글').text if law.find('법령명한글') is not None else None
                            break
                    print(f"[DEBUG] law_id: {law_id}, law_name_kor: {law_name_kor}")
                    if not law_id or not law_name_kor:
                        print(f"[DEBUG] 법령ID 또는 법령명한글 없음, continue")
                        continue
                    # 2. lawService로 본문 요청
                    detail_params = {
                        'OC': OC,
                        'target': 'law',
                        'type': 'XML',
                        'ID': law_id
                    }
                    detail_resp = requests.get(detail_url, params=detail_params, timeout=60)
                    print(f"[DEBUG] lawService status: {detail_resp.status_code}")
                    detail_root = ET.fromstring(detail_resp.text)
                    upper_law_text = ''
                    for node in detail_root.iter():
                        if node.tag == '조문내용' and node.text and node.text.strip():
                            content = re.sub(r'<[^>]+>', '', node.text)
                            content = content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').strip()
                            upper_law_text += content + '\n'
                        elif node.tag == '항내용' and node.text and node.text.strip():
                            content = re.sub(r'<[^>]+>', '', node.text)
                            content = content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').strip()
                            upper_law_text += '    ' + content + '\n'
                        elif node.tag == '호내용' and node.text and node.text.strip():
                            content = re.sub(r'<[^>]+>', '', node.text)
                            content = content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').strip()
                            upper_law_text += '        ' + content + '\n'
                    print(f"[DEBUG] upper_law_text length: {len(upper_law_text)}")
                    if not upper_law_text.strip():
                        print(f"[DEBUG] upper_law_text가 비어 있음, continue")
                        continue
                    # 상위법령 검토 결과 추가
                    doc.add_heading(f'상위 법령명: {upper_law_name}', level=2)
                    doc.add_paragraph('(아래는 상위 법령 전체 조문 중 조례와 직접적으로 관련 있는 조문만 발췌/요약한 내용입니다.)')
                    doc.add_paragraph(upper_law_text[:2000])  # 본문 일부도 워드에 기록
                    # Gemini API로 위반 여부 분석
                    if 'geminiApiKey' in request.form:
                        try:
                            genai.configure(api_key=request.form['geminiApiKey'])
                            model = genai.GenerativeModel('gemini-1.5-flash')
                            prompt = (
                                f'아래는 상위 법령({upper_law_name})의 전체 조문과 내가 업로드한 조례의 전체 내용이야.\n'
                                '---상위 법령---\n'
                                f'{upper_law_text}\n'
                                '---내 조례---\n'
                                f'{pdf_text[:2000]}\n'
                                '---\n'
                                '상위 법령 전체 조문 중에서, 내가 업로드한 조례와 직접적으로 관련 있는 조문(또는 위반 가능성이 있는 조문)만 발췌해서 요약해줘. 반드시 한글로 답변해줘.\n'
                                '1. [법령우위의 원칙 위반 여부]\n'
                                '- 조례가 상위 법령의 내용과 직접적으로 충돌하거나 위배되는지\n'
                                '- 상위 법령의 취지나 목적을 해치는지\n'
                                '- 상위 법령이 금지하는 행위를 허용하거나, 의무화하는 행위를 면제하는지\n\n'
                                '2. [법률 유보의 원칙 위반 여부]\n'
                                '- 주민의 권리를 제한하거나 의무를 부과하는 내용이 있는지\n'
                                '- 상위 법령에서 위임받지 않은 권한을 행사하는지\n'
                                '- 상위 법령의 위임 범위를 초과하는지\n\n'
                                '3. [실무적 검토 포인트]\n'
                                '- 조례의 집행 과정에서 발생할 수 있는 문제점\n'
                                '- 상위 법령과의 관계에서 주의해야 할 사항\n'
                                '- 개선이 필요한 부분과 그 방향성\n'
                            )
                            print(f"[DEBUG] Gemini 프롬프트 길이: {len(prompt)}")
                            response = model.generate_content(prompt)
                            print(f"[DEBUG] Gemini 응답: {getattr(response, 'text', None)}")
                            if response and hasattr(response, 'text') and response.text:
                                clean_gemini = re.sub(r'[\*#`>\-]+', '', response.text)
                                for line in clean_gemini.split('\n'):
                                    if line.strip():
                                        doc.add_paragraph(line.strip())
                            else:
                                doc.add_paragraph('Gemini API 응단이 비어있음 또는 None입니다.')
                        except Exception as e:
                            print(f"Gemini API 오류: {e}")
                            doc.add_paragraph(f"상위법령 위반 여부 분석 중 오류가 발생했습니다: {str(e)}")

                except Exception as e:
                    print(f"상위법령 검토 중 오류 발생: {e}")
                    continue

    return doc

def is_valid_law_name(name):
    """
    법령명이 유효한지 검사하는 함수 (더 엄격하게 개선)
    """
    name_clean = name.strip().replace(' ', '').replace('「','').replace('」','').replace('[','').replace(']','').replace('(','').replace(')','').replace('<','').replace('>','').replace('"','').replace('"','').lower()
    # 너무 짧은 이름 제외 (4글자 미만 제외)
    if len(name_clean) < 4:
        print(f'[DEBUG] 너무 짧아서 제외: {name}')
        return False
    # 반드시 '법', '시행령', '시행규칙'으로 끝나야 함
    if not (name_clean.endswith('법') or name_clean.endswith('시행령') or name_clean.endswith('시행규칙')):
        print(f'[DEBUG] 법/시행령/시행규칙으로 끝나지 않아 제외: {name}')
        return False
    # 단독 불용어 제외
    if name_clean in {'기본법', '시행령', '시행규칙'}:
        print(f'[DEBUG] 단독 불용어로 제외: {name}')
        return False
    # "등 관련 법" 및 유사 표현 제외
    if '등관련법' in name_clean or '관련법' in name_clean or '관련법령' in name_clean:
        print(f'[DEBUG] 관련법 표현으로 제외: {name}')
        return False
    # "및 관련 시행령" 및 유사 표현 제외
    if '및관련시행령' in name_clean or '및시행령' in name_clean or '및시행규칙' in name_clean:
        print(f'[DEBUG] 관련 시행령/규칙 표현으로 제외: {name}')
        return False
    # "령과의 법" 및 유사 표현 제외
    if '령과의법' in name_clean or '령과법' in name_clean or '령과의' in name_clean:
        print(f'[DEBUG] 령과의 법 표현으로 제외: {name}')
        return False
    # 불용어 목록
    invalid_terms = {
        '자치입법', '조례', '규칙', '지침', '내규', '예규', '훈령', '적법', '입법', 
        '상위법', '위법', '합법', '불법', '방법', '헌법상', '헌법적', '법적', 
        '법률적', '법령상', '법률상', '법률', '법령', '법', '규정', '조항', 
        '조문', '규범', '원칙', '기준', '사항', '관련법', '관련 법', '관련법령', 
        '관련 법령', '이러한표현이모호하여법', '법령우위의원칙위반여부'
    }
    for invalid in invalid_terms:
        if name_clean == invalid.replace(' ', '').lower():
            print(f'[DEBUG] 불용어로 제외: {name}')
            return False
    # 숫자, 영문, 특수문자 포함된 경우 제외
    if not all('\uac00' <= ch <= '\ud7a3' for ch in name_clean):
        print(f'[DEBUG] 한글 이외 문자 포함되어 제외: {name}')
        return False
    return True

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port) 